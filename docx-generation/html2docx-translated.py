'''
Author: Vaibhav Singh
Editor: Pranav Dalvi
'''

# Import libraries
from bs4 import BeautifulSoup as bs
from docx import Document
from docx.shared import Inches, Cm
import re
import glob
import os
import argparse

document = Document() # Function call to initialize a document instance to which I write
font = document.styles['Normal'].font
font.name = 'Shobhika Regular' # Change font property

# def convert2docx(body_children, body, prev_element, document): # Helps me to extract tagwise text or media, and add every component to a Microsoft word document using python-docx library
#     for ocr_carea in prev_element.find_all('div'):
#         for ocr_par in ocr_carea.find_all('p'):
#             paragraph = document.add_paragraph()
#             paragraph_text = '' 
#             for ocr_line in ocr_par.find_all('span', class_ = 'ocr_line'):
#                 for ocr_word in ocr_line.find_all('span', class_ = 'ocrx_word'):
#                     paragraph_text = paragraph_text + ' ' + ocr_word.get_text()
#             paragraph.text = paragraph_text
#     document.add_page_break() # I add a page break after processing one html file

# def convert2docx(body_children,body, prev_element, document):
#     for page_tag in body.find_all("div"):
#         for paragraph_tag in page_tag.find_all('p'):
#             for word_tag in paragraph_tag.find_all('span'):
#                 paragraph = document.add_paragraph()
#                 paragraph.text = word_tag.get_text()
#         document.add_page_break()

def convert2docx(body_children,body, prev_element, document):
    for paragraph_tag in body.find_all('p'):
        for word_tag in paragraph_tag.find_all('span'):
            paragraph = document.add_paragraph()
            paragraph.text = word_tag.get_text()
    document.add_page_break()

def find_children(body): # This function helps me find the number of immediate children the <body> tag has
    temp = body.find_next().find_next()
    body_children = 1
    all_siblings = len(temp.find_next_siblings())
    body_children += all_siblings
    prev_element = body.find_next()
    print(body_children)
    convert2docx(body_children,body, prev_element, document) # Call to convert2docx() function

def initialize(soup):
    body = soup.find("body") # Helps me find <body> tag, to extract all its content
    find_children(body) # Call to find_children() function

def main(repo_name):
    # args = parser.parse_args()
    # repo_link = "https://github.com/UdaanContentForLogging/" + repo_name
    # clone_link = repo_link[repo_link.index("github"):]+".git"
    # os.system("git clone https://code-with-Aniket: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@" + clone_link)

    doc_name = f"./{repo_name}/CorrectorOutput/p-*.html"
    file_paths = glob.glob(doc_name)
    sorted_file_paths = sorted(file_paths)
    print('initializing document(.docx) generation')
    for file in sorted_file_paths:
        print(f'...processing file: {file}')
        filename = open(file, 'r', encoding="utf-8") # Instead of passing html file directly to beautifulsoup, I open html file and pass the filehandle into Beautiful Soup.
        soup = bs(filename, 'html.parser')
        sections = document.sections
        for section in sections:
            section.top_margin = Cm(1.27)
            section.bottom_margin = Cm(1.27)
            section.left_margin = Cm(1.27)
            section.right_margin = Cm(1.27)
        # soup = bs(soup.prettify(), 'html.parser') 
        # soup.prettify()
        # print(soup) # I am using bs4 html parser to parse all pages of html content, the html content takes the form of a tree
        initialize(soup) # Call to initialize() function

    document.save(repo_name + '.docx') # Save document
    print('document(.docx) generation successful')

if __name__ == '__main__':
    # repo_names = ["Vasant_Lad_ocr_indic_pdf","Vinod_Verma-Sixteen_Minutes_to_a_Better_9-to-5_1999", "SUSHEE_1", "Central_Council_for_Research_in_Ayurveda_and_Siddha","Josep_G.V.R_and_Central_Council_for_Research_in_Ayurveda_and_Siddha_India","Manisha_Kshirsagar-Ayurveda_A_Quick_Reference_Handbook_Lotus_Press_WI_2014"]
    repo_names = ["Software_Wikipedia"]

    for repo_name in repo_names:
        main(repo_name)

# to run this python script: 1) Give name of the Repo in repo_names = ["___"]
                    #2) in Terminal --- change environment as test-environment and change directory as cd pranav/docx-generation>
                    #3) then type ( python html2docx-translated.py ) and enter