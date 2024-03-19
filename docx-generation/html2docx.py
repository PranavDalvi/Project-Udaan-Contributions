# Import libraries
from bs4 import BeautifulSoup as bs
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Cm
from docx.shared import Pt
from docx.shared import RGBColor
import re
import glob

document = Document() # Function call to initialize a document instance to which I write
font = document.styles['Normal'].font
font.name = 'Shobhika Regular' # Change font property
sections = document.sections
for section in sections:
    section.top_margin = Cm(1.27)
    section.bottom_margin = Cm(1.27)
    section.left_margin = Cm(1.27)
    section.right_margin = Cm(1.27)

def convert2docx_pre_edit(div_children, prev_element, document): # Helps me to extract tagwise text or media, and add every component to a Microsoft word document using python-docx library
    p_text = ''
    for element in range(div_children):
        if prev_element.name == 'p' or prev_element.name == 'table' or prev_element.name == 'ul' or prev_element.name == 'ol':
            current_element = prev_element.find_next_sibling()
        else:
            current_element = prev_element.find_next()

        if current_element.name == 'img':
            if current_element.has_attr('src'):
                image_link = current_element['src']
                document.add_picture(image_link, width = Inches(3))
            # document.add_heading(current_element, level=3)
            prev_element = current_element
        elif current_element.name == 'p':
            p_children = len(current_element.find_all())
            prev = current_element
            for element in range(p_children):
                    current = prev.find_next()
                    p_text += current.get_text(strip=True)
                    p_text += ' '
                    prev = current
            document.add_paragraph(p_text)
            p_text = ''
            prev_element = current_element
        elif current_element.name == 'span':
            document.add_heading(current_element.get_text(), level=2)
            prev_element = current_element
        elif current_element.name == 'ul':
            list_elements = len(current_element.find_next().find_next_siblings()) + 1
            prev_list_element =  current_element
            for li in range(list_elements):
                current_list_element = prev_list_element.find_next()
                document.add_paragraph(current_list_element.get_text(strip = True), style = 'List Bullet')
                prev_list_element = current_list_element
            prev_element = current_element
        elif current_element.name == 'ol':
            list_elements = len(current_element.find_next().find_next_siblings()) + 1
            prev_list_element =  current_element
            for li in range(list_elements):
                current_list_element = prev_list_element.find_next()
                document.add_paragraph(current_list_element.get_text(strip = True), style = 'List Number')
                prev_list_element = current_list_element
            prev_element = current_element
        elif current_element.name == 'table':
            column_temp = current_element.find("td")
            row_temp = current_element.find_next()
            table_rows = len(row_temp.find_next_siblings()) + 1
            tr_tags = current_element.find_all('tr')
            max_row_td_length = 0
            max_row_th_length = 0
            for tr in tr_tags:
                td = tr.find_all("td")
                img = tr.find_all("img")
                if len(td) + len(img) > max_row_td_length:
                    max_row_td_length = len(td) + len(img)
            for tr in tr_tags:
                th = tr.find_all("th")
                if len(th) > max_row_th_length:
                    max_row_th_length = len(th)
            if max_row_td_length >= max_row_th_length:
                table_columns = max_row_td_length
            else:
                table_columns = max_row_th_length 
            if table_columns == 0:
                continue
            print("table rows = ", table_rows)
            prev_row_element = current_element
            # table_columns = len(column_temp.find_next_siblings()) + 1
            print("table columns = ", table_columns)
            table = document.add_table(rows = table_rows, cols = table_columns)
            table.style = 'Table Grid'
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            # table.allow_autofit = True
            for row_element in range(table_rows):
                current_row_element = prev_row_element.find_next()
                table_columns = len(current_row_element.find_all())
                prev_column_element = current_row_element
                # if current_row_element.find_next().name == 'th':
                #     hdr_cells = table.rows[row_element].cells
                # else:
                row_cells = table.rows[row_element].cells
                for column_element in range(table_columns):
                    column_content = ''
                    current_column_element = prev_column_element.find_next()
                    if current_column_element.name == 'img':
                        if current_column_element.has_attr('src'):
                            image_link = current_column_element['src']
                            para = row_cells[column_element].paragraphs[0]
                            run = para.add_run()
                            run.add_picture(image_link, width = Inches(3))
                    else:
                        column_content = current_column_element.get_text(strip=True)
                        # print(current_column_element
                        if current_column_element.name == 'th':
                            row_cells[column_element].paragraphs[0].add_run(column_content).bold = True
                        else:
                            row_cells[column_element].text = column_content
                        # print(column_content)
                    prev_column_element = current_column_element
                prev_row_element = current_column_element
            prev_element = current_element
    document.add_page_break() # I add a page break after processing one html file


def find_children_pre_edit(div): # This function helps me find the number of immediate children the <body> tag has
    temp = div.find_next()
    div_children = 1
    all_siblings = len(temp.find_next_siblings())
    div_children += all_siblings
    prev_element = div
    convert2docx_pre_edit(div_children, prev_element, document) # Call to convert2docx() function

def initialize_pre_edit(soup):
    div = soup.find("div", attrs={"class": "ocr_page"})
    find_children_pre_edit(div) # Call to find_children() function

def main(repo_name):
    print('initializing document(.docx) generation')
    for file in sorted(glob.glob(f"./{repo_name}/CorrectorOutput/p-*.html")):
        print(f'...processing file: {file}')
        filename = open(file, 'r', encoding="utf8") # Instead of passing html file directly to beautifulsoup, I open html file and pass the filehandle into Beautiful Soup.
        soup = bs(filename, 'html.parser') # I am using bs4 html parser to parse all pages of html content, the html content takes the form of a tree
        initialize_pre_edit(soup) # Call to initialize() function

    document.save('./' + repo_name + '.docx') # Save document
    print('document(.docx) generation successful')

if __name__ == '__main__':
    repo_names = ["Agni_parvathalu"]

    for repo_name in repo_names:
        main(repo_name)

    
